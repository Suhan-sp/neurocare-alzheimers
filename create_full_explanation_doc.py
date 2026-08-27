"""
Script to generate 'Fullexplanation.doc' Word document for NeuroCare Project & Research Paper Defense.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    title_color = RGBColor(15, 23, 42)      # Dark Navy / Slate
    heading_color = RGBColor(13, 148, 136)  # Accent Teal
    subheading_color = RGBColor(79, 70, 229)# Accent Indigo
    body_color = RGBColor(51, 65, 85)      # Slate Dark Text

    # Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = body_color

    # TITLE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("NEUROCARE AI: COMPREHENSIVE PROJECT & RESEARCH PAPER EXPLANATION GUIDE")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = title_color

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Master Codebase Mapping, Mathematical Formulations, Experimental Basis & Panel Defense\nTeam 36 — Sai Vidya Institute of Technology, Bengaluru\n")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 1: EXECUTIVE SUMMARY
    h1 = doc.add_heading("1. Executive Project Summary & Paper Abstract", level=1)
    h1.runs[0].font.color.rgb = heading_color

    p = doc.add_paragraph()
    p.add_run("NeuroCare ").bold = True
    p.add_run("is a non-invasive, multimodal clinical decision-support platform engineered for the early detection and risk stratification of Alzheimer's Disease (AD). Traditional diagnostic frameworks rely on cost-prohibitive, invasive neuroimaging (Positron Emission Tomography - PET, structural Magnetic Resonance Imaging - MRI) or lumbar spinal taps (cerebrospinal fluid - CSF analysis). Standardized paper-based cognitive screening tests like the Mini-Mental State Examination (MMSE) are geographically accessible but suffer from subjective examiner scoring variability and practice effects.\n\n")
    p.add_run("To solve these single-modality vulnerabilities, NeuroCare synthesizes three distinct non-invasive biomedical channels: ")
    p.add_run("(1) Neuropsychological Cognitive Metrics ").bold = True
    p.add_run("(processed from the OASIS longitudinal dataset, N = 373), ")
    p.add_run("(2) 19-Channel Continuous Scalp Electroencephalogram (EEG) Signals ").bold = True
    p.add_run("(processed from OpenNeuro ds004504, N = 88 human subjects, 1,320 clinical epochs) enhanced by a Computer-Vision EEG Report Digitizer, and ")
    p.add_run("(3) Speech Acoustic Biomarkers ").bold = True
    p.add_run("(derived from N = 591 clinical audio recordings). Probability outputs from individual modality classifiers are fused using a decision-level Sequential Least Squares Programming (SLSQP) soft-voting ensemble layer, accompanied by SHapley Additive exPlanations (SHAP) for patient-level explainability.\n\n")

    p.add_run("Experimental 5-Fold Stratified Cross-Validation demonstrates that the multimodal ensemble achieves a ").bold = True
    p.add_run("96.85% validation accuracy ").bold = True
    p.add_run("and an ")
    p.add_run("ROC-AUC of 0.9882").bold = True
    p.add_run(", outperforming all standalone single-modality classifiers and existing literature baselines.")

    # SECTION 2: CODEBASE FILE-TO-PAPER SECTION MAP
    h2 = doc.add_heading("2. Master Codebase File-to-Paper Section Mapping", level=1)
    h2.runs[0].font.color.rgb = heading_color

    p = doc.add_paragraph("This section provides an authoritative mapping between every section of the research paper (paper.tex) and the exact Python source code files, classes, functions, and saved model binary artifacts in the project workspace:\n")

    # Table of Code Map
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Paper Section", "Source Code File Path", "Core Classes / Functions", "Saved Artifacts / Datasets"]
    for i, t in enumerate(hdr_titles):
        hdr_cells[i].text = t
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "0F766E") # Dark Teal

    code_map_data = [
        ("Section III-A\nCognitive Pipeline", "preprocessing/cognitive_preprocessor.py\nmodels/ml_models.py\ntrain_cognitive.py", "CognitivePreprocessor.preprocess()\nMLModelSuite.train_and_evaluate_all()", "saved_models/cognitive_preprocessor.pkl\nsaved_models/best_cognitive_model.pkl\n(OASIS Dataset N=373)"),
        ("Section III-B\n19-Ch EEG Pipeline", "preprocessing/eeg_preprocessor.py\nfeature_extraction/eeg_features.py\ntrain_eeg.py", "EEGPreprocessor.bandpass_filter()\nextract_multichannel_eeg_features()\nWelch PSD & Hjorth Parameters", "saved_models/eeg_preprocessor.pkl\nsaved_models/best_eeg_model.pkl\n(OpenNeuro ds004504 N=88, 1320 epochs)"),
        ("Section III-C\nVision EEG Digitizer", "services/eeg_digitizer.py", "EEGReportDigitizer.digitize_report()\nPyMuPDF, EasyOCR, OpenCV crop,\nskimage.skeletonize (1-pixel peak tracking)", "static/plots/digitized_eeg_waveform.png\n(Confidence Score: 95.0%)"),
        ("Section III-D\nSpeech Acoustic Pipeline", "preprocessing/speech_preprocessor.py\nfeature_extraction/speech_features.py\ntrain_speech.py", "SpeechPreprocessor.extract_features()\nLibrosa 80+ acoustic features\nRandom Forest Ensemble (200 trees)", "saved_models/speech_preprocessor.pkl\nsaved_models/best_speech_model.pkl\n(Cookie Theft N=591 recordings)"),
        ("Section III-E\nSLSQP Ensemble Fusion", "ensemble.py", "MultimodalEnsemble.optimize_weights()\npredict_ensemble() using scipy.optimize.minimize", "Fusion Weights: [0.34, 0.33, 0.33]\nRisk Levels: Low, Moderate, High"),
        ("Section III-F\nSHAP XAI Layer", "utils/explainability.py", "SHAPExplainer.explain_prediction()\nTreeExplainer / LinearExplainer", "static/plots/shap_waterfall.png\nstatic/plots/shap_summary.png"),
        ("Section IV\nImplementation & Backend", "app.py\ndatabase.py\nprod_server.py", "Flask 3.0.0, Waitress WSGI (Port 5000)\ninit_db(), save_patient_report()\nget_patients_grouped_by_user()", "neurocare.db (SQLite Database)\nUser Roles: Patient, Doctor, Admin"),
        ("Section V\nResults & Evaluation", "evaluate_ensemble.py\nutils/visualization.py", "evaluate_system()\n5-Fold Stratified Cross-Validation\nROC-AUC & Confusion Matrix plotting", "Table II (Metrics: 96.85% Acc)\nTable III (Literature Benchmark)")
    ]

    for row_idx, row_data in enumerate(code_map_data):
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(row_cells[i], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 3: MATHEMATICAL FORMULATIONS & EQUATION DERIVATIONS
    h3 = doc.add_heading("3. Mathematical Formulations & Equation Derivations", level=1)
    h3.runs[0].font.color.rgb = heading_color

    p = doc.add_paragraph("This section breaks down all 12 mathematical equations specified in the research paper, detailing their exact physical/statistical meaning and their implementation location in code:\n")

    eqs = [
        ("Equation (1): Cognitive Z-Score Normalization",
         "x' = (x - μ) / σ",
         "Standardizes Oasis demographic, cognitive, and volumetric clinical variables to zero mean and unit variance. Implemented in preprocessing/cognitive_preprocessor.py using StandardScaler.fit_transform()."),
        
        ("Equation (2): Regularized Logistic Regression Sigmoid Log-Loss",
         "P(AD_cog | x_cog) = 1 / (1 + exp(-(w^T x' + b)))",
         "Converts weighted cognitive feature vectors into calibrated output probabilities P_cognitive in [0, 1]. Implemented in models/ml_models.py using scikit-learn's LogisticRegression with L2 regularization penalty."),

        ("Equation (3): 4th-Order Zero-Phase Butterworth Bandpass Filter Transfer Function",
         "|H(f)|^2 = 1 / (1 + (f / f_c)^(2N))  [Range: 0.5 Hz to 45.0 Hz]",
         "Attenuates low-frequency DC offsets (< 0.5 Hz) and high-frequency muscle/blink artifacts (> 45 Hz). Implemented in preprocessing/eeg_preprocessor.py using scipy.signal.butter(4, [0.5, 45], btype='band') and scipy.signal.filtfilt()."),

        ("Equation (4): Spectral Slowing Ratio (Theta/Alpha Power Ratio)",
         "R_(theta/alpha) = P_theta / P_alpha",
         "Quantifies cortical spectral slowing—a hallmark pathophysiological indicator of Alzheimer's Disease where high-frequency alpha power (8–12 Hz) shifts toward lower-frequency theta power (4–8 Hz). Implemented in feature_extraction/eeg_features.py using Welch Power Spectral Density (PSD) integration."),

        ("Equations (5), (6), (7): Hjorth Time-Domain Signal Parameters",
         "Activity = Var(y(t))\nMobility = sqrt(Var(y'(t)) / Var(y(t)))\nComplexity = Mobility(y'(t)) / Mobility(y(t))",
         "Measures signal power variance (Activity), mean frequency representation (Mobility), and change in signal frequency complexity (Complexity). Implemented in feature_extraction/eeg_features.py."),

        ("Equation (8): Speech Pitch Local Jitter",
         "Jitter(local) = [ (1 / (N-1)) * Σ |T_i - T_(i+1)| ] / [ (1 / N) * Σ T_i ]",
         "Quantifies frame-to-frame fundamental frequency (F0) pitch period perturbations caused by vocal cord micro-tremors and motor control degradation in early dementia. Implemented in feature_extraction/speech_features.py using Librosa."),

        ("Equations (9), (10), (11): SLSQP Soft-Voting Multimodal Ensemble Fusion & Weight Optimization",
         "P_ensemble = Σ (W_i * P_i)  subject to  Σ W_i = 1, W_i ≥ 0\nOptimal Weights: W_cognitive = 0.34, W_eeg = 0.33, W_speech = 0.33",
         "Decision-level soft-voting probability fusion optimized via Sequential Least Squares Programming (SLSQP) by minimizing validation log-loss. Implemented in ensemble.py using scipy.optimize.minimize(method='SLSQP')."),

        ("Equation (12): SHAP (SHapley Additive exPlanations) Feature Attribution Value",
         "φ_j = Σ_{S ⊆ F \\ {j}} [ |S|!(|F| - |S| - 1)! / |F|! ] * [ f(S ∪ {j}) - f(S) ]",
         "Cooperative game theory formulation assigning a fair marginal contribution value φ_j to each patient feature j across all feature subsets S. Implemented in utils/explainability.py using shap.TreeExplainer and shap.LinearExplainer.")
    ]

    for eq_title, eq_formula, eq_desc in eqs:
        doc.add_heading(eq_title, level=2).runs[0].font.color.rgb = subheading_color
        p_eq = doc.add_paragraph()
        p_eq.add_run(eq_formula).bold = True
        p_eq.runs[0].font.color.rgb = RGBColor(15, 118, 110)
        p_desc = doc.add_paragraph(eq_desc)
        p_desc.paragraph_format.space_after = Pt(8)

    # SECTION 4: EXPERIMENTAL RESULTS & LITERATURE COMPARISON
    h4 = doc.add_heading("4. Experimental Results & Literature Benchmark", level=1)
    h4.runs[0].font.color.rgb = heading_color

    doc.add_heading("Table II: Performance Metrics of the Proposed NeuroCare System (5-Fold Cross Validation)", level=2).runs[0].font.color.rgb = subheading_color

    table2 = doc.add_table(rows=1, cols=6)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2_cells = table2.rows[0].cells
    hdr2_titles = ["Modality / Architecture", "Machine Learning Model", "Validation Acc (%)", "ROC-AUC", "Sensitivity (%)", "Specificity (%)"]
    for i, t in enumerate(hdr2_titles):
        hdr2_cells[i].text = t
        hdr2_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr2_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr2_cells[i], "4F46E5") # Dark Indigo

    table2_data = [
        ("Cognitive Data", "Regularized Logistic Regression", "94.64%", "0.9677", "93.85%", "95.20%"),
        ("19-Channel EEG Signal", "XGBoost Classifier", "77.50%", "0.8203", "76.10%", "78.60%"),
        ("Speech Acoustic", "Random Forest Ensemble (200 Trees)", "93.02%", "0.9775", "91.50%", "92.70%"),
        ("NeuroCare Multimodal", "SLSQP Soft-Voting Ensemble", "96.85%", "0.9882", "96.20%", "97.35%")
    ]

    for row_idx, row_data in enumerate(table2_data):
        row_cells = table2.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if "Multimodal" in row_data[0]:
                row_cells[i].paragraphs[0].runs[0].font.bold = True
                set_cell_background(row_cells[i], "E0E7FF")
            else:
                set_cell_background(row_cells[i], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.add_heading("Table III: Comparison with Existing State-of-the-Art Literature", level=2).runs[0].font.color.rgb = subheading_color

    table3 = doc.add_table(rows=1, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3_cells = table3.rows[0].cells
    hdr3_titles = ["Study / Reference", "Methodology & Modalities Used", "Accuracy (%)"]
    for i, t in enumerate(hdr3_titles):
        hdr3_cells[i].text = t
        hdr3_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr3_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr3_cells[i], "0F766E")

    table3_data = [
        ("Marcus et al. (2010)", "Demographic SVM (Cognitive & MRI features only)", "88.50%"),
        ("Jeong et al. (2004)", "Spectral Relative Power (Standalone EEG spectral analysis)", "76.20%"),
        ("Cummins et al. (2020)", "MFCC + Acoustic SVM (Standalone speech acoustics)", "86.40%"),
        ("Achmad et al. (2023)", "Feature Concatenation (Early feature-level fusion)", "91.20%"),
        ("NeuroCare (Proposed)", "SLSQP Decision Soft Ensemble + SHAP (Multimodal)", "96.85%")
    ]

    for row_idx, row_data in enumerate(table3_data):
        row_cells = table3.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if "NeuroCare" in row_data[0]:
                row_cells[i].paragraphs[0].runs[0].font.bold = True
                set_cell_background(row_cells[i], "CCFBF1")
            else:
                set_cell_background(row_cells[i], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 5: COMPREHENSIVE PANEL Q&A VIVA DEFENSE
    h5 = doc.add_heading("5. Comprehensive Panel Viva & Conference Defense (Q&A)", level=1)
    h5.runs[0].font.color.rgb = heading_color

    qas = [
        ("Q1: Why is standalone EEG accuracy 77.50% while Cognitive is 94.64% and Speech is 93.02%?",
         "Answer: 19-channel scalp EEG recordings are naturally subject to scalp impedance, ocular blinks, and muscle noise. However, EEG provides high temporal resolution insight into cortical 'spectral slowing' (theta/alpha power ratio) prior to overt behavioral symptoms. When combined in our SLSQP soft-voting ensemble, EEG provides critical neurodynamic information that elevates the overall system accuracy to 96.85%."),

        ("Q2: How does the Vision EEG Report Digitizer work when raw .edf digital signal files are unavailable?",
         "Answer: The vision digitizer (services/eeg_digitizer.py) follows a 5-step pipeline: (1) PyMuPDF renders PDF pages to 150 DPI matrices, (2) EasyOCR extracts clinical text findings ('alpha rhythm', 'theta slowing'), (3) OpenCV performs adaptive Gaussian thresholding and grid region cropping, (4) skimage.morphology.skeletonize thins binary curves into 1-pixel skeletons, and (5) median column-wise displacement tracking extracts continuous time-series arrays y(t) resampled at 500 Hz across 19 standard channels with 95.0% extraction confidence."),

        ("Q3: Why did you choose Decision-Level Soft Voting (SLSQP) instead of Early Feature Concatenation?",
         "Answer: Early feature concatenation (as used in Achmad et al. 2023, 91.20% accuracy) suffers from the curse of dimensionality and feature scale distortion when combining tabular scores, high-frequency EEG arrays, and acoustic audio features. Decision-level soft voting independently trains optimized classifiers on each modality's native feature space, converting them to calibrated probabilities [0, 1] before SLSQP weight optimization (0.34, 0.33, 0.33), achieving superior resilience (96.85% accuracy)."),

        ("Q4: Are the machine learning models retrained every time a user submits an assessment on the web app?",
         "Answer: No. The machine learning models are pre-trained ONCE offline (train_cognitive.py, train_eeg.py, train_speech.py) using 5-fold cross-validation. Their optimized binary weights and scaling objects are saved in the saved_models/ directory as .pkl files. When a user submits an assessment on http://localhost:5000, predict.py instantly loads the .pkl files via joblib.load() and performs sub-millisecond real-time inference."),

        ("Q5: How does the system handle missing optional modalities (e.g., if a clinic uploads only Cognitive + Voice without EEG)?",
         "Answer: In ensemble.py, the MultimodalEnsemble.predict_ensemble() method dynamically normalizes weights across ONLY user-provided active input channels (e.g., reallocating weights between Cognitive and Speech) without injecting dummy zero-padding or fake fallback data, maintaining strict medical evaluation integrity."),

        ("Q6: How is clinical data persisted and secured in the backend?",
         "Answer: Patient evaluation reports, channel power spectral densities, EasyOCR findings, and diagnosis timestamps are stored in an embedded SQLite database (neurocare.db) using parameterized SQL queries to prevent SQL injection. Role-Based Access Control (RBAC) enforces strict session boundaries separating Patient Portal (/patient/login), Doctor Portal (/doctor/login), and Admin Dashboard (/admin)."),

        ("Q7: How did you calculate the confidence score of the system prediction?",
         "Answer: In ensemble.py, prediction confidence is calculated as the absolute distance from the uncertain 0.5 decision boundary: Confidence = |P_ensemble - 0.5| * 2.0 * 100%, bounded between 50.0% and 99.9%."),

        ("Q8: How does SHAP provide clinical interpretability for physicians?",
         "Answer: SHAP (SHapley Additive exPlanations) uses cooperative game theory (Equation 12) to compute individual feature contribution values φ_j for each patient. On the web result page (result.html), doctors see a visual waterfall plot showing exactly how much each biomarker (e.g., MMSE score of 22 or Theta slowing) shifted the patient's risk score above or below the baseline population mean.")
    ]

    for q, a in qas:
        doc.add_heading(q, level=2).runs[0].font.color.rgb = subheading_color
        doc.add_paragraph(a).paragraph_format.space_after = Pt(8)

    # Save document
    filename = "Fullexplanation.doc"
    doc.save(filename)
    print(f"Successfully generated {filename}!")

if __name__ == '__main__':
    create_document()
